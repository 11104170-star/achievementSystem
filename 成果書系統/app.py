import streamlit as st
import pandas as pd
from collections import Counter
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from io import BytesIO
import os

# =========================
# 頁面設定
# =========================

st.set_page_config(
    page_title="成果書生成器",
    page_icon="📄",
    layout="wide"
)

st.title("📄 茶道社成果書生成器")

# =========================
# 上傳檔案
# =========================

template_file = st.file_uploader(
    "上傳 Word 模板",
    type=["docx"]
)

excel_file = st.file_uploader(
    "上傳問卷 Excel / CSV",
    type=["xlsx", "csv"]
)

# =========================
# 活動資訊
# =========================

st.header("活動資訊")

fill_date = st.text_input("填寫日期")

activity_name = st.text_input("活動名稱")

activity_date = st.text_input("活動日期")

activity_place = st.text_input("活動地點")

activity_people = st.text_input("參加人數")

activity_leader = st.text_input("活動負責人")

phone = st.text_input("連絡電話")

activity_review = st.text_area("活動檢討")

teacher_comment = st.text_area("指導老師評語")

photo1_desc = st.text_input("照片1說明")

photo2_desc = st.text_input("照片2說明")

photo3_desc = st.text_input("照片3說明")

# =========================
# 上傳照片
# =========================

st.header("上傳照片")

flow_photo = st.file_uploader(
    "活動流程照片",
    type=["jpg", "jpeg", "png"]
)

group_photo = st.file_uploader(
    "大合照",
    type=["jpg", "jpeg", "png"]
)

photo1 = st.file_uploader(
    "照片1",
    type=["jpg", "jpeg", "png"]
)

photo2 = st.file_uploader(
    "照片2",
    type=["jpg", "jpeg", "png"]
)

photo3 = st.file_uploader(
    "照片3",
    type=["jpg", "jpeg", "png"]
)

# =========================
# 中文數字
# =========================

chinese_numbers = {
    0: "零",
    1: "一",
    2: "二",
    3: "三",
    4: "四",
    5: "五",
    6: "六",
    7: "七",
    8: "八",
    9: "九",
    10: "十",
    11: "十一",
    12: "十二",
    13: "十三",
    14: "十四",
    15: "十五",
    16: "十六",
    17: "十七",
    18: "十八",
    19: "十九",
    20: "二十",
}

# =========================
# 設定字型
# =========================

def set_font(paragraph):

    for run in paragraph.runs:

        run.font.name = "標楷體"

        run._element.rPr.rFonts.set(
            qn("w:eastAsia"),
            "標楷體"
        )

        run.font.size = Pt(11)

# =========================
# 替換文字
# =========================

def replace_text(doc, replace_map):

    # 正文
    for para in doc.paragraphs:

        for key, value in replace_map.items():

            if key in para.text:

                para.text = para.text.replace(
                    key,
                    value
                )

                set_font(para)

    # 表格
    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for para in cell.paragraphs:

                    for key, value in replace_map.items():

                        if key in para.text:

                            para.text = para.text.replace(
                                key,
                                value
                            )

                            set_font(para)

# =========================
# 插入圖片
# =========================

def insert_images(doc, image_map):

    # 正文
    for para in doc.paragraphs:

        for key, image_file in image_map.items():

            if key in para.text:

                para.text = ""

                if image_file is not None:

                    run = para.add_run()

                    run.add_picture(
                        image_file,
                        width=Inches(2.5)
                    )

    # 表格
    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for para in cell.paragraphs:

                    for key, image_file in image_map.items():

                        if key in para.text:

                            para.text = ""

                            if image_file is not None:

                                run = para.add_run()

                                run.add_picture(
                                    image_file,
                                    width=Inches(2.5)
                                )

# =========================
# 生成成果書
# =========================

if st.button("生成成果書"):

    if template_file is None:

        st.error("請上傳 Word 模板")
        st.stop()

    if excel_file is None:

        st.error("請上傳問卷 Excel")
        st.stop()

    # =========================
    # 讀取 Excel / CSV
    # =========================

    ext = os.path.splitext(
        excel_file.name
    )[1].lower()

    try:

        if ext == ".xlsx":

            df = pd.read_excel(
                excel_file
            )

        else:

            encodings = [
                "utf-8-sig",
                "utf-8",
                "cp950",
                "big5"
            ]

            success = False

            for enc in encodings:

                try:

                    excel_file.seek(0)

                    df = pd.read_csv(
                        excel_file,
                        encoding=enc
                    )

                    success = True
                    break

                except:
                    pass

            if not success:

                st.error("CSV 編碼錯誤")
                st.stop()

    except Exception as e:

        st.error("讀取檔案失敗")
        st.exception(e)
        st.stop()

    # =========================
    # 刪除不需要欄位
    # =========================

    remove_columns = [
        "請選擇今天社課名稱",
        "學校",
        "姓名："
    ]

    for col in remove_columns:

        if col in df.columns:

            df = df.drop(columns=[col])

    # =========================
    # 問卷分析
    # =========================

    total = len(df)

    score_questions = []

    text_questions = []

    for col in df.columns:

        values = (
            df[col]
            .dropna()
            .astype(str)
        )

        is_score = True

        for v in values:

            if v.strip() not in [
                "1",
                "2",
                "3",
                "4",
                "5"
            ]:

                is_score = False
                break

        if is_score:

            score_questions.append(col)

        else:

            text_questions.append(col)

    result_text = ""

    result_text += (
        f"填寫回饋表單人數：{total}人\n\n"
    )

    # =========================
    # 評分題
    # =========================

    if len(score_questions) > 0:

        result_text += (
            "題目（1-5分，1分最低、5分最高）\n"
        )

        for i, question in enumerate(
            score_questions,
            start=1
        ):

            counter = Counter()

            values = (
                df[question]
                .dropna()
                .astype(str)
            )

            for v in values:

                counter[v.strip()] += 1

            result_text += (
                f"{i}.{question}\n"
            )

            temp = []

            for score in sorted(
                counter.keys(),
                reverse=True
            ):

                count = counter[score]

                percent = (
                    count / total * 100
                )

                count_text = chinese_numbers.get(
                    count,
                    str(count)
                )

                temp.append(
                    f"{count_text}人{score}分（{percent:.2f}%）"
                )

            result_text += (
                "、".join(temp)
            )

            result_text += "\n"

    # =========================
    # 文字題
    # =========================

    start_index = (
        len(score_questions) + 1
    )

    for idx, question in enumerate(
        text_questions,
        start=start_index
    ):

        result_text += (
            f"\n{idx}.{question}\n"
        )

        counter = Counter()

        blank = 0

        for v in df[question]:

            if pd.isna(v):

                blank += 1
                continue

            text = str(v).strip()

            if text == "":

                blank += 1

            else:

                counter[text] += 1

        for text, count in counter.items():

            result_text += (
                f"● {text}（{count}）\n"
            )

        result_text += (
            f"● 空白（{blank}）\n"
        )

    # =========================
    # 讀取 Word
    # =========================

    doc = Document(template_file)

    # =========================
    # 文字替換
    # =========================

    replace_map = {

        "{{填寫日期}}": fill_date,

        "{{活動名稱}}": activity_name,

        "{{活動日期}}": activity_date,

        "{{活動地點}}": activity_place,

        "{{參加人數}}": activity_people,

        "{{活動負責人}}": activity_leader,

        "{{連絡電話}}": phone,

        "{{活動檢討}}": activity_review,

        "{{指導老師評語}}": teacher_comment,

        "{{照片1說明}}": photo1_desc,

        "{{照片2說明}}": photo2_desc,

        "{{照片3說明}}": photo3_desc,

        "{{問卷分析結果}}": result_text,

    }

    replace_text(
        doc,
        replace_map
    )

    # =========================
    # 圖片替換
    # =========================

    image_map = {

        "{{活動流程照片}}": flow_photo,

        "{{大合照}}": group_photo,

        "{{照片1}}": photo1,

        "{{照片2}}": photo2,

        "{{照片3}}": photo3,

    }

    insert_images(
        doc,
        image_map
    )

    # =========================
    # 輸出 Word
    # =========================

    output = BytesIO()

    doc.save(output)

    output.seek(0)

    st.success("成果書生成完成！")

    st.download_button(
        label="📥 下載成果書",
        data=output,
        file_name="成果書.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
